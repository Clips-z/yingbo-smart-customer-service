"""
RAG 知识库服务 - OpenAI 兼容 API
使用 ChromaDB + 硅基流动 API 实现本地知识库检索增强生成
增强功能：递归分块策略 + Reranking 重排序
"""
import os
import site
import sys
from pathlib import Path

RAG_DIR = Path(__file__).resolve().parent
RAG_PACKAGES = RAG_DIR.parent / "tools" / "rag-py311"
site.addsitedir(str(RAG_PACKAGES))

# 在导入 chromadb 之前禁用遥测，避免 posthog capture() 兼容性告警
os.environ["CHROMA_TELEMETRY_DISABLED"] = "1"
os.environ["ANONYMIZED_TELEMETRY"] = "False"

try:
    import posthog

    posthog.disabled = True
    posthog.capture = lambda *args, **kwargs: None
except ImportError:
    pass

import re
import json
import uuid
import time
import httpx
from typing import List, Optional, Dict, Any, Tuple

import chromadb
from chromadb.config import Settings
from fastapi import FastAPI, UploadFile, File, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from openai import OpenAI

# ============================================================
# 配置管理
# ============================================================
BASE_DIR = Path(__file__).parent
CONFIG_PATH = BASE_DIR / "config.json"
DATA_DIR = BASE_DIR / "data"
CHROMA_DIR = DATA_DIR / "chroma"
UPLOAD_DIR = DATA_DIR / "uploads"

for d in [DATA_DIR, CHROMA_DIR, UPLOAD_DIR]:
    d.mkdir(parents=True, exist_ok=True)


def load_config() -> dict:
    if CONFIG_PATH.exists():
        with open(CONFIG_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}


def save_config(cfg: dict):
    with open(CONFIG_PATH, "w", encoding="utf-8") as f:
        json.dump(cfg, f, ensure_ascii=False, indent=2)


config = load_config()

# ============================================================
# ChromaDB 初始化
# ============================================================
chroma_client = chromadb.PersistentClient(
    path=str(CHROMA_DIR),
    settings=Settings(anonymized_telemetry=False),
)
collection = chroma_client.get_or_create_collection(
    name="knowledge_base",
    metadata={"hnsw:space": "cosine"}
)

# ============================================================
# OpenAI 客户端 (指向硅基流动)
# ============================================================
def get_api_key() -> str:
    api_key = config.get("siliconflow_api_key", "")
    if not api_key:
        raise HTTPException(status_code=500, detail="SiliconFlow API Key 未配置，请在管理页面设置")
    return api_key


def get_client() -> OpenAI:
    return OpenAI(
        api_key=get_api_key(),
        base_url=config.get("siliconflow_base_url", "https://api.siliconflow.cn/v1")
    )


def get_base_url() -> str:
    return config.get("siliconflow_base_url", "https://api.siliconflow.cn/v1").rstrip("/")


# ============================================================
# 文档处理
# ============================================================
def read_pdf(file_path: Path) -> str:
    from pypdf import PdfReader
    reader = PdfReader(str(file_path))
    texts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            texts.append(text.strip())
    return "\n\n".join(texts)


def read_text(file_path: Path) -> str:
    encodings = ["utf-8", "gbk", "gb2312", "latin-1"]
    for enc in encodings:
        try:
            return file_path.read_text(encoding=enc)
        except (UnicodeDecodeError, Exception):
            continue
    return ""


def read_docx(file_path: Path) -> str:
    """读取 DOCX 文件"""
    try:
        from docx import Document
        doc = Document(str(file_path))
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                # 检测标题样式
                if para.style.name.startswith('Heading'):
                    level = para.style.name.replace('Heading ', '')
                    paragraphs.append('#' * int(level) + ' ' + para.text.strip())
                else:
                    paragraphs.append(para.text.strip())
        # 也读取表格内容
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(' | '.join(cells))
        return '\n\n'.join(paragraphs)
    except ImportError:
        raise HTTPException(status_code=500, detail="缺少 python-docx 依赖，请运行: pip install python-docx")


def read_html(file_path: Path) -> str:
    """读取 HTML 文件，提取纯文本"""
    try:
        from bs4 import BeautifulSoup
        html = file_path.read_text(encoding='utf-8')
        soup = BeautifulSoup(html, 'html.parser')
        # 移除 script 和 style 标签
        for tag in soup(['script', 'style', 'nav', 'footer', 'header']):
            tag.decompose()
        return soup.get_text(separator='\n', strip=True)
    except ImportError:
        raise HTTPException(status_code=500, detail="缺少 beautifulsoup4 依赖，请运行: pip install beautifulsoup4")


def read_excel(file_path: Path) -> str:
    """读取 Excel 文件"""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(file_path), read_only=True, data_only=True)
        all_texts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            all_texts.append(f'# 工作表: {sheet_name}')
            rows = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(cell) if cell is not None else '' for cell in row]
                if any(c.strip() for c in cells):
                    rows.append(' | '.join(cells))
            all_texts.append('\n'.join(rows))
        wb.close()
        return '\n\n'.join(all_texts)
    except ImportError:
        raise HTTPException(status_code=500, detail="缺少 openpyxl 依赖，请运行: pip install openpyxl")


def read_document(file_path: Path) -> str:
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return read_pdf(file_path)
    elif suffix in (".txt", ".md", ".markdown"):
        return read_text(file_path)
    elif suffix == ".docx":
        return read_docx(file_path)
    elif suffix in (".html", ".htm"):
        return read_html(file_path)
    elif suffix in (".xlsx", ".xls"):
        return read_excel(file_path)
    elif suffix in (".json", ".csv"):
        return read_text(file_path)
    else:
        return read_text(file_path)


# ============================================================
# 递归分块策略 (Markdown 感知)
# ============================================================

# Markdown 标题正则：匹配 # ~ ###### 开头的行
MD_HEADING_RE = re.compile(r'^(#{1,6})\s+(.+)$', re.MULTILINE)

# 中文句子分隔符
SENTENCE_SPLIT_RE = re.compile(r'(?<=[。！？!?；;\n])')

# 代码块正则
CODE_BLOCK_RE = re.compile(r'(```[\s\S]*?```)', re.MULTILINE)


def split_by_markdown_headings(text: str) -> List[Tuple[str, str]]:
    """
    按 Markdown 标题切分文本，返回 (标题上下文, 正文内容) 列表。
    每个 section 包含标题层级路径作为上下文。
    """
    sections = []
    lines = text.split('\n')
    current_heading_stack = []  # [(level, title), ...]
    current_content = []

    for line in lines:
        match = MD_HEADING_RE.match(line)
        if match:
            # 保存前一个 section
            if current_content:
                heading_ctx = ' > '.join(t for _, t in current_heading_stack)
                sections.append((heading_ctx, '\n'.join(current_content).strip()))
                current_content = []

            level = len(match.group(1))
            title = match.group(2).strip()

            # 维护标题栈
            while current_heading_stack and current_heading_stack[-1][0] >= level:
                current_heading_stack.pop()
            current_heading_stack.append((level, title))
        else:
            current_content.append(line)

    # 最后一个 section
    if current_content:
        heading_ctx = ' > '.join(t for _, t in current_heading_stack)
        sections.append((heading_ctx, '\n'.join(current_content).strip()))

    return sections


def split_by_paragraphs(text: str) -> List[str]:
    """按双换行（段落）分割"""
    return [p.strip() for p in re.split(r'\n\s*\n', text) if p.strip()]


def split_by_sentences(text: str, max_length: int = 500) -> List[str]:
    """按句子分隔符分割，尽量不超过 max_length"""
    sentences = SENTENCE_SPLIT_RE.split(text)
    sentences = [s.strip() for s in sentences if s.strip()]

    result = []
    current = ""
    for sent in sentences:
        if len(current) + len(sent) <= max_length:
            current = (current + sent).strip()
        else:
            if current:
                result.append(current)
            current = sent
    if current:
        result.append(current)
    return result


def split_by_chars(text: str, chunk_size: int, overlap: int) -> List[str]:
    """按字符数硬切分（最后手段）"""
    chunks = []
    step = chunk_size - overlap
    for i in range(0, len(text), step):
        chunk = text[i:i + chunk_size]
        if chunk.strip():
            chunks.append(chunk)
        if i + chunk_size >= len(text):
            break
    return chunks


def add_overlap(chunks: List[str], overlap: int) -> List[str]:
    """为分块列表添加重叠内容"""
    if overlap <= 0 or len(chunks) <= 1:
        return chunks

    result = [chunks[0]]
    for i in range(1, len(chunks)):
        prev_tail = chunks[i - 1][-overlap:] if len(chunks[i - 1]) >= overlap else chunks[i - 1]
        result.append(prev_tail + chunks[i])
    return result


def chunk_text_recursive(
    text: str,
    chunk_size: int = 500,
    overlap: int = 100,
    min_chunk_size: int = 50,
) -> List[str]:
    """
    递归分块策略：
    1. 先保护代码块（不拆分代码块内部）
    2. 按 Markdown 标题切分
    3. 按段落（双换行）切分
    4. 按句子（。！？!?；;\n）切分
    5. 按字符数硬切分（最后手段）

    每个 chunk 会带上其所属的标题路径作为上下文。
    """
    if not text or not text.strip():
        return []

    chunk_size = chunk_size or config.get("chunk_size", 500)
    overlap = overlap or config.get("chunk_overlap", 100)
    min_chunk_size = max(min_chunk_size, 20)

    final_chunks = []

    # Step 1: 保护代码块，先提取出来
    code_blocks = {}
    def replace_code_block(match):
        key = f"__CODE_BLOCK_{len(code_blocks)}__"
        code_blocks[key] = match.group(0)
        return key

    protected_text = CODE_BLOCK_RE.sub(replace_code_block, text)

    # Step 2: 按 Markdown 标题切分
    sections = split_by_markdown_headings(protected_text)

    for heading_ctx, section_text in sections:
        if not section_text:
            continue

        # 如果 section 本身就很短，直接作为一个 chunk
        if len(section_text) <= chunk_size:
            prefix = f"[{heading_ctx}]\n" if heading_ctx else ""
            final_chunks.append(prefix + section_text)
            continue

        # Step 3: 按段落切分
        paragraphs = split_by_paragraphs(section_text)

        current_chunk = ""
        for para in paragraphs:
            # 恢复代码块占位符
            for key, code in code_blocks.items():
                para = para.replace(key, code)
                current_chunk = current_chunk.replace(key, code)

            if len(current_chunk) + len(para) + 2 <= chunk_size:
                current_chunk = (current_chunk + "\n\n" + para).strip() if current_chunk else para
            else:
                if current_chunk:
                    prefix = f"[{heading_ctx}]\n" if heading_ctx else ""
                    final_chunks.append(prefix + current_chunk)
                    current_chunk = ""

                # 单个段落超过 chunk_size
                if len(para) > chunk_size:
                    # Step 4: 按句子切分
                    sentences = split_by_sentences(para, chunk_size)

                    sent_chunk = ""
                    for sent in sentences:
                        if len(sent_chunk) + len(sent) <= chunk_size:
                            sent_chunk = (sent_chunk + sent).strip()
                        else:
                            if sent_chunk:
                                final_chunks.append(sent_chunk)
                            # 单个句子仍然太长
                            if len(sent) > chunk_size:
                                # Step 5: 按字符硬切分
                                char_chunks = split_by_chars(sent, chunk_size, overlap)
                                final_chunks.extend(char_chunks)
                            else:
                                sent_chunk = sent
                    if sent_chunk:
                        current_chunk = sent_chunk
                else:
                    current_chunk = para

        if current_chunk:
            prefix = f"[{heading_ctx}]\n" if heading_ctx else ""
            final_chunks.append(prefix + current_chunk)

    # 如果没有 Markdown 标题结构（sections 只有一个且无标题），回退到段落+句子分块
    if len(final_chunks) == 0:
        paragraphs = split_by_paragraphs(text)
        current_chunk = ""
        for para in paragraphs:
            if len(current_chunk) + len(para) + 2 <= chunk_size:
                current_chunk = (current_chunk + "\n\n" + para).strip() if current_chunk else para
            else:
                if current_chunk:
                    final_chunks.append(current_chunk)
                if len(para) > chunk_size:
                    sentences = split_by_sentences(para, chunk_size)
                    for sent in sentences:
                        if len(sent) <= chunk_size:
                            final_chunks.append(sent)
                        else:
                            final_chunks.extend(split_by_chars(sent, chunk_size, overlap))
                else:
                    current_chunk = para
                current_chunk = "" if len(para) > chunk_size else para
        if current_chunk:
            final_chunks.append(current_chunk)

    # 过滤太短的 chunk，合并到前一个
    merged = []
    for chunk in final_chunks:
        if merged and len(chunk) < min_chunk_size:
            merged[-1] = merged[-1] + "\n" + chunk
        else:
            merged.append(chunk)

    # 添加重叠
    merged = add_overlap(merged, overlap)

    return merged


# 保留旧函数名兼容
def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
    return chunk_text_recursive(text, chunk_size, overlap)


# ============================================================
# Embedding & Reranking
# ============================================================

# 查询 embedding 缓存（LRU，避免重复查询时反复调用 API）
_embedding_cache: Dict[str, Tuple[List[float], float]] = {}
_EMBEDDING_CACHE_TTL = 300  # 5 分钟
_EMBEDDING_CACHE_MAX = 200  # 最多缓存 200 条


def _get_cached_embedding(query: str) -> List[float]:
    """带缓存的查询 embedding，避免相同查询反复调用 API"""
    now = time.time()
    cached = _embedding_cache.get(query)
    if cached and (now - cached[1]) < _EMBEDDING_CACHE_TTL:
        return cached[0]

    embedding = get_embedding([query])[0]

    # 清理过期缓存
    if len(_embedding_cache) >= _EMBEDDING_CACHE_MAX:
        expired_keys = [k for k, (_, ts) in _embedding_cache.items() if now - ts > _EMBEDDING_CACHE_TTL]
        for k in expired_keys:
            del _embedding_cache[k]
        # 如果还是太多，删掉最旧的
        if len(_embedding_cache) >= _EMBEDDING_CACHE_MAX:
            sorted_keys = sorted(_embedding_cache, key=lambda k: _embedding_cache[k][1])
            for k in sorted_keys[: len(_embedding_cache) - _EMBEDDING_CACHE_MAX + 1]:
                del _embedding_cache[k]

    _embedding_cache[query] = (embedding, now)
    return embedding


def get_embedding(texts: List[str]) -> List[List[float]]:
    """调用硅基流动 embedding API"""
    client = get_client()
    model = config.get("embedding_model", "BAAI/bge-m3")
    # 分批处理，每批最多 64 条
    all_embeddings = []
    batch_size = 64
    for i in range(0, len(texts), batch_size):
        batch = texts[i:i + batch_size]
        resp = client.embeddings.create(model=model, input=batch)
        all_embeddings.extend([item.embedding for item in resp.data])
    return all_embeddings


def rerank_documents(query: str, documents: List[str], top_n: int = 5) -> List[Tuple[str, float]]:
    """
    使用硅基流动 Reranker API 对检索结果进行重排序。
    返回 [(document, relevance_score), ...] 按相关性降序排列。

    如果 reranker 不可用，返回原始顺序。
    """
    if not documents:
        return []

    top_n = min(top_n, len(documents))
    rerank_model = config.get("rerank_model", "BAAI/bge-reranker-v2-m3")

    # 如果文档数量很少（<= top_n），且 reranker 未启用，直接返回
    use_rerank = config.get("use_rerank", True)
    if not use_rerank:
        return [(doc, 0.0) for doc in documents[:top_n]]

    try:
        api_key = get_api_key()
        base_url = get_base_url()
        url = f"{base_url}/rerank"

        # 分批处理，API 限制每次最多 100 个文档
        all_results = []
        batch_size = 100
        offset = 0

        for i in range(0, len(documents), batch_size):
            batch = documents[i:i + batch_size]
            batch_top_n = min(top_n, len(batch))

            payload = {
                "model": rerank_model,
                "query": query,
                "documents": batch,
                "top_n": batch_top_n,
                "return_documents": False,
                "max_chunks_per_doc": 1024,
            }

            with httpx.Client(timeout=30.0) as http_client:
                resp = http_client.post(
                    url,
                    json=payload,
                    headers={
                        "Authorization": f"Bearer {api_key}",
                        "Content-Type": "application/json",
                    },
                )

            if resp.status_code != 200:
                print(f"[RAG] Reranker API 返回 {resp.status_code}: {resp.text[:200]}")
                # 降级：返回原始顺序
                return [(doc, 0.0) for doc in documents[:top_n]]

            data = resp.json()
            results = data.get("results", [])

            for item in results:
                idx = item["index"] + offset
                score = item["relevance_score"]
                all_results.append((documents[idx], score))

            offset += len(batch)

        # 按分数降序排序
        all_results.sort(key=lambda x: x[1], reverse=True)
        return all_results[:top_n]

    except Exception as e:
        print(f"[RAG] Reranking 失败，降级为原始顺序: {e}")
        return [(doc, 0.0) for doc in documents[:top_n]]


# ============================================================
# 检索逻辑（向量检索 + Reranking）
# ============================================================
def retrieve_with_reranking(query: str, top_k: int = 5, candidate_multiplier: int = 3) -> Dict:
    """
    两阶段检索：
    1. 向量检索获取 top_k * candidate_multiplier 个候选
    2. Reranker 对候选重排序，取最终 top_k 个

    返回 {"query": ..., "results": [{"content": ..., "source": ..., "score": ..., "rerank_score": ...}]}
    """
    if not config.get("siliconflow_api_key"):
        raise HTTPException(status_code=400, detail="请先设置 API Key")

    # 查询预处理：去除多余空白
    query = query.strip()
    if not query:
        return {"query": query, "results": [], "total_chunks": 0}

    k = top_k or config.get("top_k", 5)
    # 候选数量 = 检索数量 * 倍数，但不超过知识库总量
    collection_count = collection.count()
    if collection_count == 0:
        return {"query": query, "results": [], "total_chunks": 0}
    candidate_n = min(k * candidate_multiplier, collection_count)

    # 可配置的相似度阈值（cosine distance < threshold 才保留）
    similarity_threshold = config.get("similarity_threshold", 0.5)
    # Reranking 最低分数（低于此分数的结果被过滤）
    min_rerank_score = config.get("min_rerank_score", 0.1)

    # Stage 1: 向量检索（带缓存）
    try:
        query_embedding = _get_cached_embedding(query)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"生成查询向量失败: {str(e)}")

    results = collection.query(
        query_embeddings=[query_embedding],
        n_results=candidate_n,
        include=["documents", "metadatas", "distances"]
    )

    docs = results.get("documents", [[]])[0]
    metas = results.get("metadatas", [[]])[0]
    dists = results.get("distances", [[]])[0]

    if not docs:
        return {"query": query, "results": [], "total_chunks": collection_count}

    # 过滤相似度太低的结果
    filtered = [
        (doc, meta, dist)
        for doc, meta, dist in zip(docs, metas, dists)
        if dist < (1 - similarity_threshold)
    ]

    if not filtered:
        return {"query": query, "results": [], "total_chunks": collection_count}

    filtered_docs = [f[0] for f in filtered]
    filtered_metas = [f[1] for f in filtered]
    filtered_dists = [f[2] for f in filtered]

    # Stage 2: Reranking
    reranked = rerank_documents(query, filtered_docs, top_n=k)

    # 组合结果
    doc_to_meta = {doc: meta for doc, meta in zip(filtered_docs, filtered_metas)}
    doc_to_dist = {doc: dist for doc, dist in zip(filtered_docs, filtered_dists)}

    final_results = []
    for doc, rerank_score in reranked:
        # 过滤 reranking 分数过低的结果
        if rerank_score < min_rerank_score:
            continue
        meta = doc_to_meta.get(doc, {})
        dist = doc_to_dist.get(doc, 1.0)
        final_results.append({
            "content": doc[:500] + "..." if len(doc) > 500 else doc,
            "full_content": doc,
            "source": meta.get("source", ""),
            "vector_score": round(1 - dist, 4),
            "rerank_score": round(rerank_score, 4),
        })

    return {"query": query, "results": final_results, "total_chunks": collection_count}


# ============================================================
# FastAPI 应用
# ============================================================
app = FastAPI(title="RAG 知识库服务", version="2.0.0")

# CORS：仅允许本地 Electron 渲染进程跨域访问
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost",
        "http://127.0.0.1",
        "http://localhost:3000",
        "http://127.0.0.1:3000",
        "http://localhost:8080",
        "http://127.0.0.1:8080",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.mount("/static", StaticFiles(directory=str(BASE_DIR / "static")), name="static")


# ============================================================
# 文件监控（自动增量更新）
# ============================================================
_watchdog_enabled = config.get("watchdog_enabled", False)
_watchdog_observer = None


def start_watchdog():
    """启动文件监控，自动检测 UPLOAD_DIR 中的文档变化"""
    global _watchdog_observer
    if not _watchdog_enabled:
        return
    try:
        from watchdog.observers import Observer
        from watchdog.events import FileSystemEventHandler

        class KnowledgeBaseHandler(FileSystemEventHandler):
            """知识库文件变更处理器"""
            def on_created(self, event):
                if not event.is_directory:
                    _handle_file_event(event.src_path, "created")

            def on_modified(self, event):
                if not event.is_directory:
                    _handle_file_event(event.src_path, "modified")

            def on_deleted(self, event):
                if not event.is_directory:
                    _handle_file_event(event.src_path, "deleted")

        def _handle_file_event(file_path, action):
            path = Path(file_path)
            suffix = path.suffix.lower()
            supported = {".pdf", ".txt", ".md", ".markdown", ".docx", ".html", ".htm", ".xlsx", ".xls", ".csv", ".json"}
            if suffix not in supported:
                return

            # 避免重复处理临时文件
            if path.name.startswith("~") or path.name.startswith("."):
                return

            print(f"[Watchdog] 文件{action}: {path.name}")

            if action in ("created", "modified"):
                try:
                    # 等待文件写入完成
                    time.sleep(1)
                    text = read_document(path)
                    if text.strip():
                        # 删除旧 chunks
                        _remove_document_chunks(path.name)
                        # 重新分块并入库
                        chunks = chunk_text_recursive(text, chunk_size, chunk_overlap)
                        _add_chunks_to_collection(chunks, path.name)
                        print(f"[Watchdog] 已更新知识库: {path.name} ({len(chunks)} 个块)")
                except Exception as e:
                    print(f"[Watchdog] 处理文件失败: {path.name} - {e}")

            elif action == "deleted":
                try:
                    _remove_document_chunks(path.name)
                    print(f"[Watchdog] 已从知识库移除: {path.name}")
                except Exception as e:
                    print(f"[Watchdog] 移除文件失败: {path.name} - {e}")

        def _remove_document_chunks(filename):
            """从 ChromaDB 中移除指定文档的所有 chunks"""
            try:
                results = collection.get(where={"source": filename})
                if results and results["ids"]:
                    collection.delete(ids=results["ids"])
            except Exception:
                pass  # 文档可能不存在

        def _add_chunks_to_collection(chunks, source_name):
            """将 chunks 添加到 ChromaDB"""
            if not chunks:
                return
            ids = [str(uuid.uuid4()) for _ in chunks]
            metadatas = [{"source": source_name, "chunk_index": i} for i in range(len(chunks))]
            collection.add(documents=chunks, ids=ids, metadatas=metadatas)

        chunk_size = config.get("chunk_size", 500)
        chunk_overlap = config.get("chunk_overlap", 50)

        _watchdog_observer = Observer()
        _watchdog_observer.schedule(
            KnowledgeBaseHandler(),
            str(UPLOAD_DIR),
            recursive=False,
        )
        _watchdog_observer.start()
        print(f"[Watchdog] 文件监控已启动，监控目录: {UPLOAD_DIR}")
    except ImportError:
        print("[Watchdog] 缺少 watchdog 依赖，文件监控未启用。请运行: pip install watchdog")
    except Exception as e:
        print(f"[Watchdog] 启动失败: {e}")


def stop_watchdog():
    """停止文件监控"""
    global _watchdog_observer
    if _watchdog_observer:
        _watchdog_observer.stop()
        _watchdog_observer.join()
        _watchdog_observer = None


@app.on_event("startup")
async def on_startup():
    start_watchdog()


@app.on_event("shutdown")
async def on_shutdown():
    stop_watchdog()


# ============================================================
# 文档管理 API
# ============================================================
@app.post("/api/upload")
async def upload_document(file: UploadFile = File(...)):
    """上传文档到知识库（文件上传）"""
    if not config.get("siliconflow_api_key"):
        raise HTTPException(status_code=400, detail="请先在管理页面设置 SiliconFlow API Key")

    # 保存文件
    file_id = str(uuid.uuid4())[:8]
    safe_name = file.filename.replace("/", "_").replace("\\", "_")
    save_path = UPLOAD_DIR / f"{file_id}_{safe_name}"

    content = await file.read()
    with open(save_path, "wb") as f:
        f.write(content)

    # 读取并分块
    text = read_document(save_path)
    if not text.strip():
        raise HTTPException(status_code=400, detail="无法从文档中提取文本内容")

    chunk_size = config.get("chunk_size", 500)
    overlap = config.get("chunk_overlap", 100)
    chunks = chunk_text_recursive(text, chunk_size, overlap)

    if not chunks:
        raise HTTPException(status_code=400, detail="文档分块后为空")

    # 生成 embedding
    try:
        embeddings = get_embedding(chunks)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"生成向量失败: {str(e)}")

    # 存入 ChromaDB
    doc_ids = [f"{file_id}_{i}" for i in range(len(chunks))]
    metadatas = [{"source": safe_name, "file_id": file_id, "chunk_index": i} for i in range(len(chunks))]

    collection.add(
        ids=doc_ids,
        documents=chunks,
        embeddings=embeddings,
        metadatas=metadatas
    )

    return {
        "success": True,
        "file_id": file_id,
        "filename": safe_name,
        "chunks": len(chunks),
        "message": f"文档已上传并分为 {len(chunks)} 个知识块（递归分块 + Reranking 已启用）"
    }


@app.post("/api/text/upload")
async def upload_text(request: Request):
    """
    上传纯文本内容到知识库（从 textarea 同步）。
    Body: {"text": "...", "filename": "knowledge_base.txt"}
    """
    if not config.get("siliconflow_api_key"):
        raise HTTPException(status_code=400, detail="请先在管理页面设置 SiliconFlow API Key")

    body = await request.json()
    text = body.get("text", "").strip()
    filename = body.get("filename", "knowledge_base.txt")

    if not text:
        raise HTTPException(status_code=400, detail="文本内容为空")

    # 先删除同名的旧文档（如果存在）
    all_data = collection.get(include=["metadatas"])
    ids_to_delete = []
    for i, meta in enumerate(all_data.get("metadatas", [])):
        if meta.get("source") == filename and meta.get("file_id", "").startswith("text_"):
            ids_to_delete.append(all_data["ids"][i])

    if ids_to_delete:
        collection.delete(ids=ids_to_delete)
        print(f"[RAG] 已删除旧文本文档 {filename} 的 {len(ids_to_delete)} 个块")

    # 分块
    file_id = f"text_{str(uuid.uuid4())[:8]}"
    chunk_size = config.get("chunk_size", 500)
    overlap = config.get("chunk_overlap", 100)
    chunks = chunk_text_recursive(text, chunk_size, overlap)

    if not chunks:
        raise HTTPException(status_code=400, detail="文本分块后为空")

    # 生成 embedding
    try:
        embeddings = get_embedding(chunks)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"生成向量失败: {str(e)}")

    # 存入 ChromaDB
    doc_ids = [f"{file_id}_{i}" for i in range(len(chunks))]
    metadatas = [{"source": filename, "file_id": file_id, "chunk_index": i} for i in range(len(chunks))]

    collection.add(
        ids=doc_ids,
        documents=chunks,
        embeddings=embeddings,
        metadatas=metadatas
    )

    return {
        "success": True,
        "file_id": file_id,
        "filename": filename,
        "chunks": len(chunks),
        "deleted_old": len(ids_to_delete),
        "message": f"文本已同步并分为 {len(chunks)} 个知识块"
    }


@app.get("/api/documents")
async def list_documents():
    """列出知识库中的所有文档"""
    all_data = collection.get(include=["metadatas"])
    doc_map = {}
    for meta in all_data.get("metadatas", []):
        fid = meta.get("file_id", "unknown")
        if fid not in doc_map:
            doc_map[fid] = {
                "file_id": fid,
                "filename": meta.get("source", "unknown"),
                "chunks": 0
            }
        doc_map[fid]["chunks"] += 1

    return {"documents": list(doc_map.values()), "total_chunks": len(all_data.get("ids", []))}


@app.delete("/api/documents/{file_id}")
async def delete_document(file_id: str):
    """删除知识库中的文档"""
    all_data = collection.get(include=["metadatas"])
    ids_to_delete = []
    for i, meta in enumerate(all_data.get("metadatas", [])):
        if meta.get("file_id") == file_id:
            ids_to_delete.append(all_data["ids"][i])

    if not ids_to_delete:
        raise HTTPException(status_code=404, detail="文档不存在")

    collection.delete(ids=ids_to_delete)

    # 删除上传的文件（仅文件上传的，文本上传的不需要删文件）
    for f in UPLOAD_DIR.glob(f"{file_id}_*"):
        f.unlink()

    return {"success": True, "deleted_chunks": len(ids_to_delete)}


@app.post("/api/config")
async def update_config(request: Request):
    """更新配置"""
    body = await request.json()
    global config
    for key, value in body.items():
        config[key] = value
    save_config(config)
    return {"success": True, "config": {k: v for k, v in config.items() if k != "siliconflow_api_key"}}


@app.get("/api/config")
async def get_config():
    """获取当前配置（隐藏 API Key）"""
    safe_config = {k: v for k, v in config.items() if k != "siliconflow_api_key"}
    safe_config["has_api_key"] = bool(config.get("siliconflow_api_key"))
    safe_config["rerank_enabled"] = config.get("use_rerank", True)
    safe_config["rerank_model"] = config.get("rerank_model", "BAAI/bge-reranker-v2-m3")
    return safe_config


@app.get("/api/search")
async def search_knowledge(query: str, top_k: int = 0):
    """
    搜索知识库（两阶段检索：向量检索 + Reranking）
    """
    result = retrieve_with_reranking(query, top_k or config.get("top_k", 5))
    return result


@app.get("/api/stats")
async def get_stats():
    """获取知识库统计"""
    count = collection.count()
    all_data = collection.get(include=["metadatas"])
    sources = set()
    for meta in all_data.get("metadatas", []):
        sources.add(meta.get("source", "unknown"))
    return {
        "total_chunks": count,
        "total_documents": len(sources),
        "has_api_key": bool(config.get("siliconflow_api_key")),
        "embedding_model": config.get("embedding_model", "BAAI/bge-m3"),
        "chat_model": config.get("chat_model", "Qwen/Qwen2.5-7B-Instruct"),
        "rerank_model": config.get("rerank_model", "BAAI/bge-reranker-v2-m3"),
        "rerank_enabled": config.get("use_rerank", True),
        "chunk_size": config.get("chunk_size", 500),
        "chunk_overlap": config.get("chunk_overlap", 100),
        "similarity_threshold": config.get("similarity_threshold", 0.5),
        "min_rerank_score": config.get("min_rerank_score", 0.1),
    }


@app.post("/api/clear")
async def clear_knowledge_base(request: Request):
    """清空全部知识库，或仅删除指定来源前缀的派生知识。"""
    try:
        body = await request.json()
    except Exception:
        body = {}
    if not isinstance(body, dict):
        raise HTTPException(status_code=400, detail="请求体必须是 JSON 对象")
    prefixes = body.get("prefixes")
    if prefixes is not None and (not isinstance(prefixes, list) or not all(isinstance(prefix, str) and prefix for prefix in prefixes)):
        raise HTTPException(status_code=400, detail="prefixes 必须是非空字符串数组")
    all_data = collection.get(include=["metadatas"])
    all_ids = all_data.get("ids", [])
    ids_to_delete = all_ids if prefixes is None else [
        all_ids[index] for index, meta in enumerate(all_data.get("metadatas", []))
        if str(meta.get("source", "")).startswith(tuple(prefixes))
    ]
    if ids_to_delete:
        collection.delete(ids=ids_to_delete)
    if prefixes is None:
        for f in UPLOAD_DIR.glob("*"):
            f.unlink()
    return {"success": True, "deleted_chunks": len(ids_to_delete)}


# ============================================================
# OpenAI 兼容 Chat API
# ============================================================
class ChatMessage(BaseModel):
    role: str
    content: str


class ChatRequest(BaseModel):
    model: str = ""
    messages: List[ChatMessage]
    stream: bool = False
    temperature: float = 0.7
    max_tokens: Optional[int] = None


def retrieve_context(query: str, top_k: int = 5) -> str:
    """检索知识库相关内容（两阶段检索 + Reranking）"""
    try:
        result = retrieve_with_reranking(query, top_k)
        docs = [r["full_content"] for r in result.get("results", [])]
        if not docs:
            return ""
        return "\n\n---\n\n".join(docs)
    except Exception as e:
        print(f"[RAG] 检索失败: {e}")
        return ""


def build_rag_messages(messages: List[ChatMessage]) -> List[ChatMessage]:
    """构建带知识库上下文的消息列表"""
    # 获取用户最后一条消息作为查询
    user_messages = [m for m in messages if m.role == "user"]
    if not user_messages:
        return messages

    last_query = user_messages[-1].content
    context = retrieve_context(last_query, config.get("top_k", 5))

    # 知识图谱增强检索
    kg_context = ""
    if config.get("kg_enabled", False):
        try:
            from kg_service import KnowledgeGraph
            kg = KnowledgeGraph(DATA_DIR)
            kg_context = kg.to_text_context(last_query, top_k=5)
        except Exception as e:
            print(f"[KG] 知识图谱检索失败: {e}")

    if not context and not kg_context:
        return messages

    # 获取原始 system prompt
    system_prompt = config.get("system_prompt", "你是一名专业的客服。")

    # 组合知识库上下文
    full_context = ""
    if context:
        full_context += f"\n【知识库参考信息（向量检索+Reranking）】\n{context}\n"
    if kg_context:
        full_context += f"\n{kg_context}\n"

    rag_prompt = f"""{system_prompt}

以下是从知识库中检索到的相关信息，请参考这些内容回答客户的问题：
{full_context}
请基于以上信息回答客户问题。如果知识库信息不足以回答，请诚实告知。"""

    # 替换 system 消息
    new_messages = [ChatMessage(role="system", content=rag_prompt)]
    new_messages.extend([m for m in messages if m.role != "system"])

    return new_messages


@app.post("/v1/chat/completions")
async def chat_completions(request: ChatRequest):
    """OpenAI 兼容的聊天接口，带 RAG 知识库检索"""
    client = get_client()
    chat_model = config.get("chat_model", "Qwen/Qwen2.5-7B-Instruct")

    # 构建 RAG 增强的消息
    rag_messages = build_rag_messages(request.messages)

    # 转换为 OpenAI 格式
    openai_messages = [{"role": m.role, "content": m.content} for m in rag_messages]

    if request.stream:
        # 流式响应
        def stream_generator():
            try:
                stream = client.chat.completions.create(
                    model=chat_model,
                    messages=openai_messages,
                    stream=True,
                    temperature=request.temperature,
                    max_tokens=request.max_tokens
                )
                for chunk in stream:
                    if chunk.choices and chunk.choices[0].delta.content:
                        delta = chunk.choices[0].delta
                        sse_data = {
                            "id": f"chatcmpl-{uuid.uuid4().hex[:8]}",
                            "object": "chat.completion.chunk",
                            "created": int(time.time()),
                            "model": chat_model,
                            "choices": [{
                                "index": 0,
                                "delta": {"content": delta.content},
                                "finish_reason": None
                            }]
                        }
                        yield f"data: {json.dumps(sse_data, ensure_ascii=False)}\n\n"
                # 结束标记
                yield f"data: {json.dumps({'id': f'chatcmpl-{uuid.uuid4().hex[:8]}', 'object': 'chat.completion.chunk', 'created': int(time.time()), 'model': chat_model, 'choices': [{'index': 0, 'delta': {}, 'finish_reason': 'stop'}]}, ensure_ascii=False)}\n\n"
                yield "data: [DONE]\n\n"
            except Exception as e:
                error_data = {"error": {"message": str(e), "type": "internal_error"}}
                yield f"data: {json.dumps(error_data, ensure_ascii=False)}\n\n"

        return StreamingResponse(stream_generator(), media_type="text/event-stream")
    else:
        # 非流式响应
        try:
            completion = client.chat.completions.create(
                model=chat_model,
                messages=openai_messages,
                stream=False,
                temperature=request.temperature,
                max_tokens=request.max_tokens
            )
            return completion.model_dump()
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))


@app.get("/v1/models")
async def list_models():
    """OpenAI 兼容的模型列表"""
    return {
        "object": "list",
        "data": [{
            "id": config.get("chat_model", "Qwen/Qwen2.5-7B-Instruct"),
            "object": "model",
            "created": int(time.time()),
            "owned_by": "siliconflow"
        }]
    }


# ============================================================
# 健康检查
# ============================================================
@app.get("/health")
async def health_check():
    """健康检查端点，供启动器检测服务是否就绪"""
    return {
        "status": "ok",
        "version": "2.0.0",
        "chunks": collection.count(),
        "rerank_enabled": config.get("use_rerank", True),
    }


# ============================================================
# Web 管理界面
# ============================================================
@app.get("/", response_class=HTMLResponse)
async def index():
    html_path = BASE_DIR / "static" / "index.html"
    return HTMLResponse(content=html_path.read_text(encoding="utf-8"))


# ============================================================
# 启动
# ============================================================
if __name__ == "__main__":
    import uvicorn
    print("=" * 60)
    print("  RAG 知识库服务 v2.0 启动中...")
    print(f"  管理页面: http://localhost:8000")
    print(f"  API 地址: http://localhost:8000/v1")
    print(f"  知识库块数: {collection.count()}")
    print(f"  分块策略: 递归 Markdown 感知分块")
    print(f"  Reranking: {'启用' if config.get('use_rerank', True) else '禁用'} ({config.get('rerank_model', 'BAAI/bge-reranker-v2-m3')})")
    print("=" * 60)
    uvicorn.run(app, host="0.0.0.0", port=8000)
